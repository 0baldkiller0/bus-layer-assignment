"""
Generate the staged Word document for PCB bus layer assignment project.
"""
import json
import os
import sys

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


# Helper functions
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def add_para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = 1.5
    return p


def add_table_from_data(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


# Load experiment results
with open('benchmark/experiment_results.json') as f:
    results = json.load(f)

# ============================================================
# Title
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PCB总线层分配与可布通性反馈优化方法')
run.bold = True
run.font.size = Pt(18)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('阶段性工作报告')
run.font.size = Pt(14)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run('2026年5月')
run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 1
# ============================================================
add_heading(doc, '1 PCB总线层分配问题', level=1)

add_para(doc, '在多层印刷电路板（PCB）设计中，总线是一组具有相同起止元件、物理走向相近的信号线集合。总线层分配的目标是为每条总线指定一个布线层，使得同层总线在物理空间上不产生冲突，同时尽量减少过孔数量、布线交叉和层间切换带来的制造成本。')

add_para(doc, '形式化地，给定总线集合B = {b1, b2, ..., bn}、可用层集合L = {l1, l2, ..., lk}以及板面障碍物集合O，层分配需要求解映射f: B -> L，使得综合代价函数C(f)最小。代价函数包含以下分量：')

items = [
    '同层冲突（conflict）：两条总线在同一层上的包围盒发生重叠，导致布线空间竞争。',
    '交叉（crossing）：两条总线的走线路径在同层交叉，需要额外绕线或过孔来避免短路。',
    '层数使用（layer usage）：实际占用的铜占用的铜层数量，层数越少越有利于制造成本控制。',
    '过孔估计（via estimate）：基于层间切换和方向转换所需的过孔数量估算。',
    '线长估计（wirelength）：总线起止点间曼哈顿距离的累加值。',
    '不可布通估计（unrouted）：与3条以上同层总线冲突的高风险总线数量。',
]
for item in items:
    add_bullet(doc, item)

add_para(doc, '综合代价为各项指标的加权和，权重分别为conflict=10、crossing=5、layer=2、via=3、wirelength=0.01、unrouted=50。其中不可布通的权重最高，因为一旦某条总线无法布通，整块板的功能将受到影响。')

# ============================================================
# 2
# ============================================================
add_heading(doc, '2 Benchmark构建与评价体系', level=1)

add_heading(doc, '2.1 合成Benchmark生成', level=2)

add_para(doc, '为系统评估不同层分配算法的性能，设计了参数化的合成benchmark生成器。每个benchmark包含板面尺寸、层数、元件集合、总线集合和障碍物集合，以JSON格式存储。总线起止位置由元件边界上的逃逸点决定，支持四种场景类型：random（总线起止位置随机分布，模拟一般性布线场景）、crossing（总线交替从板面上下方向走线，产生大量交叉冲突）、dense（总线集中在板面中心区域，产生高密度空间竞争）、parallel（总线平行排列，测试通道竞争场景）。')

add_para(doc, '测试集覆盖总线数量（5/10/20/50）、层数（2/4）和场景类型（random/crossing/dense）三个维度的组合，共生成24个benchmark实例。每个总线包含2到8个网络信号，宽度按信号数乘以0.3mm计算。')

add_heading(doc, '2.2 评价指标体系', level=2)

add_para(doc, '评价器接收benchmark JSON和层分配方案，输出六项指标及综合代价。冲突检测基于总线包围盒的AABB重叠测试：若两条同层总线的包围盒在x和y方向均有重叠，则判定为冲突。交叉检测使用线段相交的向量叉积算法，判断两条总线的直线路径是否相交。拥塞估计将板面划分为2mm*2mm栅格，统计每个栅格-层上的总线密度累积值。过孔估计考虑两类过孔：方向转换过孔（起止逃逸方向无交集时产生）和层间连接过孔（同一网络的分属不同层的总线之间需要跨层连接）。')

add_heading(doc, '2.3 真实KiCad数据导出', level=2)

add_para(doc, '系统支持从KiCad PCB设计文件导出benchmark JSON。提取流程读取.kicad_pcb和.kicad_pro文件，解析元件位置与旋转、焊盘的绝对坐标与形状尺寸、网络与焊盘的映射关系、总线起止焊盘和逃逸方向。坐标统一转换为板框左下角为原点的绝对坐标系，确保与评价器的坐标规范一致。')

# ============================================================
# 3
# ============================================================
add_heading(doc, '3 基线方法', level=1)

add_heading(doc, '3.1 Random', level=2)
add_para(doc, '随机层分配。为每条总线独立均匀随机选择一个层，作为性能下界参考。每个benchmark运行5次取最优结果，避免单次随机的偶然性。')

add_heading(doc, '3.2 Greedy', level=2)
add_para(doc, '贪心层分配。按总线宽度降序排列总线（宽度大的优先分配），逐条将总线分配到当前冲突最少的可用层。具体地，对每条待分配总线，遍历所有可用层，计算该总线与该层上已有总线的包围盒冲突数，选择冲突最少的层。时间复杂度为O(n^2 * k)，其中n为总线数，k为层数。')

add_heading(doc, '3.3 GraphColoring', level=2)
add_para(doc, '图着色启发式。以总线为节点、包围盒重叠关系为边构建冲突图，采用贪心着色策略逐个分配颜色（层）。当前实现作为启发式基线，当着色所需颜色数超过可用层数时采用取模映射，可能引入额外冲突。后续工作中需改进该方法，使其在固定层数约束下追求最小冲突着色。')

# ============================================================
# 4
# ============================================================
add_heading(doc, '4 可布通性反馈优化方法', level=1)

add_para(doc, '上述基线方法均属于单次分配策略，一旦分配完成便不再调整。反馈优化方法的核心思想是在贪心初始解的基础上，通过多轮迭代的"诊断-调整-评估"循环逐步改善分配质量。算法框架包含四个模块：失败模式识别、候选动作生成与预评分、模拟退火选择、迭代收敛控制。')

add_heading(doc, '4.1 失败模式识别', level=2)
add_para(doc, '失败检测器分析当前层分配中的结构性问题，识别三类失败模式。')
add_para(doc, '第一类为高冲突模式。某条总线与同层2条及以上总线的包围盒发生重叠，指示该总线应当被移出当前层。判定条件为冲突数达到最大冲突数的50%以上。')
add_para(doc, '第二类为层过载模式。某层上的总线数量超过各层平均值的1.5倍，且存在实际冲突。该模式指示该层需要疏散部分总线到其他层。')
add_para(doc, '第三类为孤立冲突模式。两条总线在同层发生唯一冲突（各自的冲突次数均不超过1）。此类冲突的解决成本最低，通常只需将其中一条移至其他层。')
add_para(doc, '所有检测到的失败模式按严重程度降序排列，其中涉及的总线被标记为优先调整对象。')

add_heading(doc, '4.2 候选动作生成与预评分', level=2)
add_para(doc, '针对优先调整总线，生成两类候选调整动作。move操作将一条总线从当前层移至目标层。swap操作将两条不同层上的总线互换层，适用于双方互换后均可改善的情况。')
add_para(doc, '每个候选动作在提交完整评价之前，先通过O(n)快速冲突估计进行预评分。具体做法是计算候选总线在目标层上会与多少条已有总线产生包围盒冲突，以新旧冲突数的差值作为评分。负分表示改善，正分表示恶化。swap操作的评分为两条总线各自冲突变化量的和。候选按预评分升序排列，仅取前5名进入完整评价。这一策略将单轮候选评估的复杂度从O(n^2 * m)降低到O(n * m + n * 5)，其中m为候选总数。')

add_heading(doc, '4.3 模拟退火选择', level=2)
add_para(doc, '完整评估后，对最优候选采用模拟退火机制决定是否接受。若候选使代价严格降低（delta < 0），直接接受。若代价升高（delta > 0），以概率exp(-delta / T)接受，其中T为温度参数。初始温度设为5.0，每轮乘以冷却系数0.92。该机制允许算法在早期阶段接受一定幅度的劣化解以跳出局部最优，随着温度下降逐渐退化为纯贪心接受策略。')

add_heading(doc, '4.4 迭代与收敛控制', level=2)
add_para(doc, '算法从贪心初始解出发，迭代执行上述诊断-生成-选择-调整流程。设置最大迭代轮数30，早停耐心值8（连续8轮无改善则终止）。记录每轮的代价、冲突数、执行动作、候选评估数和接受状态，用于分析收敛行为和调试。')

# ============================================================
# 5
# ============================================================
add_heading(doc, '5 实验结果', level=1)

add_heading(doc, '5.1 主实验：四种方法对比', level=2)

add_para(doc, '在24个合成benchmark上运行Random、Greedy、GraphColoring和FeedbackOpt四种方法，综合代价对比结果如表1和表2所示。')

bench_order = sorted([k for k in results if 'layers' in k and 'buses' in k
                      and 'example' not in k and 'ablation' not in k])

for layer_count in [2, 4]:
    layer_benches = [k for k in bench_order if f'layers{layer_count}' in k]
    if not layer_benches:
        continue

    add_para(doc, f'表1.{layer_count}：{layer_count}层板综合代价对比', bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER)

    headers = ['Benchmark', 'Random', 'Greedy', 'GraphColoring', 'FeedbackOpt', '改善(%)']
    rows = []
    for k in layer_benches:
        r = results[k]
        short = k.replace('buses', 'B').replace(f'layers{layer_count}_', '_L')
        row = [short]
        for method in ['Random', 'Greedy', 'GraphColoring', 'Feedback']:
            if method in r and 'metrics' in r[method]:
                row.append(f'{r[method]["metrics"]["cost"]:.1f}')
            else:
                row.append('-')
        # improvement
        g = r.get('Greedy', {}).get('metrics', {}).get('cost', 0)
        f = r.get('Feedback', {}).get('metrics', {}).get('cost', 0)
        imp = (g - f) / g * 100 if g > 0 else 0
        row.append(f'{imp:+.1f}')
        rows.append(row)

    add_table_from_data(doc, headers, rows)
    doc.add_paragraph()

add_para(doc, '从实验结果可以观察到以下规律。')

add_para(doc, '第一，在层数受限（2层）的场景下，FeedbackOpt相对于Greedy的改善幅度较为显著。在buses10_layers2_dense上改善11.4%，buses20_layers2_dense上改善8.7%。2层板的可行解空间小，初始贪心解的优化余量大，反馈迭代能够有效识别并调整高冲突总线。')

add_para(doc, '第二，在层数充裕（4层）的场景下，改善效果与场景密度密切相关。buses20_layers4_dense（高密度场景）上FeedbackOpt实现18.8%的改善，这是所有测试中改善幅度最大的。而在随机场景中（如buses10_layers4_random），Greedy本身已接近最优，反馈优化不再有改善空间。')

add_para(doc, '第三，GraphColoring在多数场景下表现不如Greedy。这是因为当前图着色实现采用取模映射策略，当着色数超过可用层数时会重新引入冲突。')

add_para(doc, '第四，大规模实例（50条总线）上，FeedbackOpt仍能保持1.2%到9.3%的改善，表明该方法具有一定的规模适应性。')

add_heading(doc, '5.2 冲突数分析', level=2)

add_para(doc, '冲突数是层分配质量最直接的指标。表2给出各方法在代表性benchmark上的冲突数对比。')

conflict_benches = ['buses20_layers2_dense', 'buses20_layers4_dense',
                    'buses50_layers2_dense', 'buses50_layers4_dense']
headers = ['Benchmark', 'Greedy', 'FeedbackOpt', 'GraphColoring']
rows = []
for k in conflict_benches:
    if k not in results:
        continue
    r = results[k]
    short = k.replace('buses', 'B').replace('layers', 'L').replace('_', ' ')
    row = [short]
    for method in ['Greedy', 'Feedback', 'GraphColoring']:
        if method in r and 'metrics' in r[method]:
            row.append(str(r[method]['metrics']['conflict_count']))
        else:
            row.append('-')
    rows.append(row)

add_table_from_data(doc, headers, rows)
doc.add_paragraph()

add_para(doc, 'FeedbackOpt在部分场景中虽然综合代价降低，但冲突数并不总是同时减少。这是因为反馈优化策略可能通过增加少量冲突来换取不可布通总线数的大幅减少，而不可布通指标的权重（50）远高于冲突权重（10）。这也说明，单一指标的对比不足以评价层分配方法的优劣，需要综合代价函数作为统一衡量标准。')

add_heading(doc, '5.3 消融实验', level=2)

add_para(doc, '为验证反馈优化方法各模块的独立贡献，在三个代表性benchmark上进行了消融实验。消融配置包括：A1_Full为完整方法（对照组）；A2_NoFailure去掉失败识别模块，改为随机选择总线尝试调整；A3_MoveOnly仅使用move策略，去掉swap操作；A4_ConflictOnly仅保留高冲突检测，去掉层过载和孤立冲突识别；A5_RandomInit以随机初始解替代贪心初始解。')

abl_data = {}
for abl_file in ['benchmark/synthetic/buses20_layers4_dense_ablation.json',
                 'benchmark/synthetic/buses20_layers2_crossing_ablation.json',
                 'benchmark/synthetic/buses50_layers4_dense_ablation.json']:
    if os.path.exists(abl_file):
        with open(abl_file) as f:
            abl_data[os.path.basename(abl_file).replace('_ablation.json', '')] = json.load(f)

for bench_name, abl in abl_data.items():
    short = bench_name.replace('buses', 'B').replace('layers', 'L').replace('_', ' ')
    add_para(doc, f'表3：消融实验结果 - {short}', bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER)

    headers = ['变体', '综合代价', '冲突数', '耗时(s)', '迭代轮数']
    rows = []
    for vname in sorted(abl.keys(), key=lambda k: abl[k]['metrics']['cost']):
        m = abl[vname]['metrics']
        rows.append([
            vname,
            f'{m["cost"]:.1f}',
            str(m['conflict_count']),
            f'{abl[vname]["time"]:.3f}',
            str(abl[vname]['iterations'])
        ])
    add_table_from_data(doc, headers, rows)
    doc.add_paragraph()

add_para(doc, '消融实验的结果揭示了以下现象。')

add_para(doc, '第一，A2（去掉失败识别）在部分场景中与A1持平甚至略优。这表明当前失败识别模块尚未充分证明其独立贡献。失败识别的设计目标是聚焦问题区域以提高搜索效率，但在当前实现中，随机选择策略有时也能覆盖到关键总线，两者效果差异不显著。后续可通过引入更细粒度的优先级评分或多轮失败历史跟踪来增强该模块的区分度。')

add_para(doc, '第二，A3（仅move）和A4（仅高冲突检测）在复杂场景中代价高于A1，说明多策略组合和多类型失败模式识别确实在发挥作用。swap操作在需要同时调整两条总线层位的场景中有其不可替代的价值。')

add_para(doc, '第三，A5（随机初始）的性能明显劣于A1，验证了贪心初始解对反馈优化的重要性。好的初始解能够减少迭代收敛所需的轮数，并提高最终解的质量。')

# ============================================================
# 6
# ============================================================
add_heading(doc, '6 系统架构', level=1)

add_para(doc, '当前系统包含以下模块，形成了从数据生成到实验评价的完整工具链。')

add_bullet(doc, 'generator.py：合成benchmark生成器，支持random/crossing/dense/parallel四种场景。')
add_bullet(doc, 'evaluator.py：可布通性评价器，计算六项指标和综合代价。')
add_bullet(doc, 'solvers.py：基线求解器，包含Random、Greedy、GraphColoring和OptimalSearch。')
add_bullet(doc, 'feedback_solver.py：反馈优化求解器，包含FailureDetector、AdjustmentGenerator和模拟退火选择。')
add_bullet(doc, 'exporter.py：从KiCad PCB文件导出benchmark JSON，保留焊盘级精度。')
add_bullet(doc, 'ablation.py：消融实验脚本，支持模块消融和参数敏感性分析。')
add_bullet(doc, 'visualize.py / viz_real_pcb.py：可视化模块，支持板面视图和对比图。')
add_bullet(doc, 'pcb_exporter.py / freerouting_workflow.py / routed_parser.py：Freerouting验证链路（初步实现）。')

# ============================================================
# 7
# ============================================================
add_heading(doc, '7 当前局限与后续工作', level=1)

add_para(doc, '当前工作存在以下需要解决的问题。')

add_para(doc, '第一，消融实验中A2（无失败识别）有时与完整方法持平，说明失败识别模块的有效性尚需进一步论证。后续可通过引入更细粒度的优先级评分或多轮失败历史跟踪来增强该模块的区分度。')

add_para(doc, '第二，GraphColoring基线的取模映射策略使其性能劣于Greedy，与严格图着色方法的定位不符。后续应改为在固定层数约束下的精确着色，或明确标注为启发式方法。')

add_para(doc, '第三，评价器中存在总线ID与列表下标隐式绑定的假设，当从真实KiCad文件导出的benchmark经过过滤或排序后可能不满足该假设，需改为基于字典的显式映射。')

add_para(doc, '第四，Freerouting验证链路的路由结果解析器仅通过网络ID判断布通状态，未实现连通性验证，存在假阳性风险。后续需改为基于几何接触关系的连通图验证。')

add_para(doc, '后续工作将沿以下方向推进。一是完善评价闭环：统一坐标系统，修复拥塞计算，实现基于连通图的布通性验证。二是强化实验论证：将FeedbackOpt纳入统一实验runner，在更广泛的参数空间上验证方法有效性，并封装本科Yan-style算法作为额外baseline。三是引入学习增强：在规则反馈优化框架稳定的基础上，探索使用轻量图神经网络替代手工预评分函数，提升候选排序的准确性。')

# ============================================================
# Save
# ============================================================
output_path = 'doc/阶段性报告_PCB总线层分配_2026-05.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'Saved: {output_path}')
print(f'File size: {os.path.getsize(output_path) / 1024:.1f} KB')
